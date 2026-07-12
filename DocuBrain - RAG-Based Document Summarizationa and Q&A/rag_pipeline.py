from __future__ import annotations

import re
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from uuid import uuid4

import pandas as pd
import pdfplumber
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

try:
    from langchain_ollama import ChatOllama
except ImportError:  # pragma: no cover - optional until Ollama support is installed
    ChatOllama = None

try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:  # pragma: no cover - optional until local embeddings are installed
    HuggingFaceEmbeddings = None

try:
    import pypdfium2 as pdfium
except ImportError:  # pragma: no cover - optional runtime dependency path
    pdfium = None

try:
    import pytesseract
except ImportError:  # pragma: no cover - optional runtime dependency path
    pytesseract = None


@dataclass
class RetrievalMetrics:
    top_k_used: int
    average_relevance: float
    precision_at_k: float
    quality_label: str
    confidence_score: float


@dataclass
class RetrievedChunk:
    document: Document
    combined_score: float
    semantic_score: float
    keyword_score: float


@dataclass
class ChunkDisplay:
    source: str
    location: str
    relevance_score: float
    highlighted_excerpt: str
    supporting_sentences: list[str]


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _split_sentences(text: str) -> list[str]:
    normalized = _normalize_text(text)
    if not normalized:
        return []
    sentences = re.split(r"(?<=[.!?])\s+", normalized)
    cleaned_sentences = []
    for sentence in sentences:
        cleaned = re.sub(r"^[^\w(]+", "", sentence.strip())
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        if len(cleaned) < 12:
            continue
        if not re.search(r"[A-Za-z0-9]", cleaned):
            continue
        cleaned_sentences.append(cleaned)
    return cleaned_sentences


def _keyword_terms(question: str) -> list[str]:
    stop_words = {
        "what",
        "when",
        "where",
        "which",
        "who",
        "why",
        "how",
        "is",
        "are",
        "was",
        "were",
        "the",
        "in",
        "on",
        "of",
        "for",
        "to",
        "a",
        "an",
        "and",
        "or",
        "did",
        "does",
        "do",
        "tell",
        "me",
        "about",
        "please",
        "from",
        "with",
        "into",
        "over",
        "each",
        "expand",
    }
    return [
        token
        for token in re.findall(r"[a-zA-Z0-9]+", question.lower())
        if len(token) > 2 and token not in stop_words
    ]


def _keyword_overlap(question: str, text: str) -> float:
    terms = _keyword_terms(question)
    if not terms:
        return 0.0

    haystack = text.lower()
    hits = sum(1 for term in terms if term in haystack)
    return hits / len(terms)


def _highlight_terms(text: str, terms: list[str]) -> str:
    highlighted = text
    for term in sorted(set(terms), key=len, reverse=True):
        pattern = re.compile(rf"(?i)\b({re.escape(term)})\b")
        highlighted = pattern.sub(r"**\1**", highlighted)
    return highlighted


def _merge_page_text(pypdf_text: str, plumber_text: str) -> str:
    normalized_pypdf = _normalize_text(pypdf_text)
    normalized_plumber = _normalize_text(plumber_text)

    if normalized_pypdf and normalized_plumber:
        if normalized_pypdf == normalized_plumber:
            return normalized_pypdf
        return (
            normalized_plumber
            if len(normalized_plumber) >= len(normalized_pypdf)
            else normalized_pypdf
        )

    return normalized_plumber or normalized_pypdf


def _is_extraction_weak(text: str) -> bool:
    normalized = _normalize_text(text)
    if len(normalized) < 40:
        return True

    alpha_numeric = re.findall(r"[A-Za-z0-9]", normalized)
    word_like = re.findall(r"\b[A-Za-z0-9]{2,}\b", normalized)
    return len(alpha_numeric) < 25 or len(word_like) < 8


def _ocr_page(file_path: Path, page_index: int) -> str:
    if pdfium is None or pytesseract is None:
        return ""

    pdf = pdfium.PdfDocument(str(file_path))
    try:
        page = pdf.get_page(page_index)
        try:
            image = page.render(scale=2.0).to_pil()
            ocr_text = pytesseract.image_to_string(image)
            return _normalize_text(ocr_text)
        finally:
            page.close()
    finally:
        pdf.close()


def extract_pdf_documents(file_path: str | Path) -> list[Document]:
    file_path = Path(file_path)
    reader = PdfReader(str(file_path))
    documents: list[Document] = []

    with pdfplumber.open(str(file_path)) as plumber_pdf:
        for page_index, page in enumerate(reader.pages):
            pypdf_text = page.extract_text() or ""
            plumber_text = ""

            if page_index < len(plumber_pdf.pages):
                plumber_text = plumber_pdf.pages[page_index].extract_text() or ""

            merged_text = _merge_page_text(pypdf_text, plumber_text)
            used_ocr = False

            # If extracted text looks too weak, attempt OCR recovery for scanned/image PDFs.
            if _is_extraction_weak(merged_text):
                ocr_text = _ocr_page(file_path, page_index)
                if ocr_text and not _is_extraction_weak(ocr_text):
                    merged_text = ocr_text
                    used_ocr = True

            if not merged_text:
                continue

            documents.append(
                Document(
                    page_content=merged_text,
                    metadata={
                        "page": page_index + 1,
                        "source": file_path.name,
                        "document_id": f"{file_path.name}_p{page_index + 1}",
                        "ocr_used": used_ocr,
                    },
                )
            )

    return documents


def extract_docx_documents(file_path: str | Path) -> list[Document]:
    file_path = Path(file_path)
    docx = DocxDocument(str(file_path))
    blocks = []

    for paragraph in docx.paragraphs:
        text = _normalize_text(paragraph.text)
        if text:
            blocks.append(text)

    for table_index, table in enumerate(docx.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [_normalize_text(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            blocks.append(f"Table {table_index}\n" + "\n".join(rows))

    text = "\n\n".join(blocks).strip()
    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": file_path.name,
                "section": "Document",
                "document_id": f"{file_path.name}_doc",
            },
        )
    ]


def extract_excel_documents(file_path: str | Path) -> list[Document]:
    file_path = Path(file_path)
    documents: list[Document] = []
    sheets = pd.read_excel(str(file_path), sheet_name=None, dtype=str, keep_default_na=False)

    for sheet_name, dataframe in sheets.items():
        if dataframe.empty:
            continue

        dataframe = dataframe.dropna(how="all").dropna(axis=1, how="all")
        if dataframe.empty:
            continue

        table_text = dataframe.to_string(index=False)
        text = _normalize_text(f"Sheet: {sheet_name}\n{table_text}")
        if not text:
            continue

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": file_path.name,
                    "sheet": str(sheet_name),
                    "document_id": f"{file_path.name}_sheet_{sheet_name}",
                },
            )
        )

    return documents


def extract_text_documents(file_path: str | Path) -> list[Document]:
    file_path = Path(file_path)

    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            text = file_path.read_text(encoding=encoding)
            break
        except UnicodeDecodeError:
            text = ""

    text = _normalize_text(text)
    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": file_path.name,
                "section": "Text",
                "document_id": f"{file_path.name}_text",
            },
        )
    ]


def extract_csv_documents(file_path: str | Path) -> list[Document]:
    file_path = Path(file_path)

    try:
        dataframe = pd.read_csv(str(file_path), dtype=str, keep_default_na=False)
        text = dataframe.to_string(index=False)
    except Exception:
        text = file_path.read_text(encoding="utf-8", errors="ignore")

    text = _normalize_text(text)
    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": file_path.name,
                "section": "CSV",
                "document_id": f"{file_path.name}_csv",
            },
        )
    ]


def extract_documents(file_path: str | Path) -> list[Document]:
    file_path = Path(file_path)
    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return extract_pdf_documents(file_path)
    if extension == ".docx":
        return extract_docx_documents(file_path)
    if extension == ".xlsx":
        return extract_excel_documents(file_path)
    if extension == ".txt":
        return extract_text_documents(file_path)
    if extension == ".csv":
        return extract_csv_documents(file_path)

    raise ValueError(f"Unsupported file type: {extension or 'unknown'}")


def split_documents(documents: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1200,
        chunk_overlap=180,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    split_docs = splitter.split_documents(documents)
    chunk_counters: dict[str, int] = {}

    for chunk in split_docs:
        source = str(chunk.metadata.get("source", "unknown"))
        page = chunk.metadata.get("page")
        sheet = chunk.metadata.get("sheet")
        section = chunk.metadata.get("section")
        location = f"p{page}" if page else f"sheet_{sheet}" if sheet else str(section or "document")
        safe_location = re.sub(r"[^A-Za-z0-9_-]+", "_", location)
        counter_key = f"{source}:{safe_location}"
        chunk_index = chunk_counters.get(counter_key, 0) + 1
        chunk_counters[counter_key] = chunk_index
        chunk.metadata["location"] = location
        chunk.metadata["chunk_id"] = f"{source}_{safe_location}_c{chunk_index}"

    return split_docs


def get_embeddings(model_name: str | None = None):
    provider = os.getenv("EMBEDDING_PROVIDER", "openai").lower().strip()

    if provider == "local":
        if HuggingFaceEmbeddings is None:
            raise ImportError(
                "Local embeddings require langchain-huggingface and sentence-transformers. "
                "Run: pip install -r requirements.txt"
            )

        return HuggingFaceEmbeddings(
            model_name=model_name
            or os.getenv("LOCAL_EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
        )

    if provider == "openai":
        return OpenAIEmbeddings(
            model=model_name or os.getenv("OPENAI_EMBEDDING_MODEL", "text-embedding-3-small")
        )

    raise ValueError("EMBEDDING_PROVIDER must be either 'openai' or 'local'.")


def build_vectorstore(
    documents: list[Document],
    embeddings: OpenAIEmbeddings,
    persist_root: str | Path,
) -> tuple[Chroma, str]:
    collection_name = f"pdf-rag-{uuid4().hex[:12]}"
    store = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory=str(persist_root),
        collection_name=collection_name,
    )
    return store, collection_name


def build_corpus_from_files(file_paths: list[str | Path]) -> list[Document]:
    page_documents: list[Document] = []
    for file_path in file_paths:
        page_documents.extend(extract_documents(file_path))
    return split_documents(page_documents)


def describe_quality(average_relevance: float, precision_at_k: float) -> str:
    score = (average_relevance * 0.7) + (precision_at_k * 0.3)

    if score >= 0.82:
        return "High"
    if score >= 0.62:
        return "Medium"
    return "Low"


def build_confidence_score(average_relevance: float, precision_at_k: float) -> float:
    confidence = (average_relevance * 0.65) + (precision_at_k * 0.35)
    return round(max(0.0, min(confidence, 1.0)), 2)


def retrieve_context(
    vectorstore: Chroma,
    question: str,
    top_k: int,
    min_relevance: float,
) -> tuple[list[RetrievedChunk], RetrievalMetrics]:
    candidate_count = max(top_k * 3, 12)
    pairs = vectorstore.similarity_search_with_score(question, k=candidate_count)

    rescored: list[RetrievedChunk] = []
    for doc, distance in pairs:
        semantic_score = 1 / (1 + max(distance, 0))
        keyword_score = _keyword_overlap(question, doc.page_content)
        combined_score = (semantic_score * 0.75) + (keyword_score * 0.25)
        rescored.append(
            RetrievedChunk(
                document=doc,
                combined_score=combined_score,
                semantic_score=semantic_score,
                keyword_score=keyword_score,
            )
        )

    rescored.sort(key=lambda item: item.combined_score, reverse=True)
    top_ranked = rescored[:top_k]

    filtered_chunks = [
        chunk
        for chunk in top_ranked
        if chunk.combined_score >= min_relevance or chunk.keyword_score >= 0.5
    ]

    if not filtered_chunks and top_ranked:
        best_chunk = top_ranked[0]
        if best_chunk.combined_score >= (min_relevance * 0.7) or best_chunk.keyword_score > 0:
            filtered_chunks = [best_chunk]

    scores = [chunk.combined_score for chunk in filtered_chunks]
    average_relevance = sum(scores) / len(scores) if scores else 0.0
    precision_at_k = len(filtered_chunks) / len(top_ranked) if top_ranked else 0.0
    quality_label = describe_quality(average_relevance, precision_at_k)
    confidence_score = build_confidence_score(average_relevance, precision_at_k)

    return filtered_chunks, RetrievalMetrics(
        top_k_used=top_k,
        average_relevance=average_relevance,
        precision_at_k=precision_at_k,
        quality_label=quality_label,
        confidence_score=confidence_score,
    )


def format_document_location(doc: Document) -> str:
    if doc.metadata.get("page"):
        return f"Page {doc.metadata.get('page')}"
    if doc.metadata.get("sheet"):
        return f"Sheet {doc.metadata.get('sheet')}"
    if doc.metadata.get("section"):
        return str(doc.metadata.get("section"))
    return "Document"


def format_context(documents: Iterable[Document]) -> str:
    blocks = []
    for index, doc in enumerate(documents, start=1):
        blocks.append(
            f"[Chunk {index} | Source {doc.metadata.get('source', 'Unknown')} | {format_document_location(doc)}]\n"
            f"{doc.page_content}"
        )
    return "\n\n".join(blocks)


def build_llm(model_name: str | None = None, temperature: float = 0.0):
    provider = os.getenv("LLM_PROVIDER", "openai").lower().strip()

    if provider == "ollama":
        if ChatOllama is None:
            raise ImportError(
                "Ollama support requires langchain-ollama. Run: pip install -r requirements.txt"
            )

        return ChatOllama(
            model=model_name or os.getenv("OLLAMA_MODEL", "llama3.1"),
            temperature=temperature,
        )

    if provider == "openai":
        return ChatOpenAI(
            model=model_name or os.getenv("OPENAI_MODEL", "gpt-4"),
            temperature=temperature,
        )

    openai_compatible_providers = {
        "groq": {
            "api_key": os.getenv("GROQ_API_KEY"),
            "model": os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile"),
            "base_url": os.getenv("GROQ_BASE_URL", "https://api.groq.com/openai/v1"),
        },
        "minimax": {
            "api_key": os.getenv("MINIMAX_API_KEY"),
            "model": os.getenv("MINIMAX_MODEL", "MiniMax-M2.5"),
            "base_url": os.getenv("MINIMAX_BASE_URL", "https://api.minimax.io/v1"),
        },
        "nvidia": {
            "api_key": os.getenv("NVIDIA_API_KEY"),
            "model": os.getenv("NVIDIA_MODEL", "meta/llama-3.1-70b-instruct"),
            "base_url": os.getenv("NVIDIA_BASE_URL", "https://integrate.api.nvidia.com/v1"),
        },
    }

    if provider in openai_compatible_providers:
        config = openai_compatible_providers[provider]
        if not config["api_key"] or config["api_key"].startswith("your_"):
            raise ValueError(f"{provider.upper()}_API_KEY is required when LLM_PROVIDER={provider}.")

        return ChatOpenAI(
            api_key=config["api_key"],
            base_url=config["base_url"],
            model=model_name or config["model"],
            temperature=temperature,
        )

    raise ValueError("LLM_PROVIDER must be one of: openai, ollama, groq, minimax, nvidia.")


def format_chat_history(chat_history: list[dict[str, str]], limit: int = 6) -> str:
    if not chat_history:
        return "No prior conversation."

    turns = chat_history[-limit:]
    lines = []
    for message in turns:
        role = message.get("role", "user").capitalize()
        content = message.get("content", "").strip()
        if content:
            lines.append(f"{role}: {content}")
    return "\n".join(lines) if lines else "No prior conversation."


def contextualize_question(
    llm: ChatOpenAI,
    question: str,
    chat_history: list[dict[str, str]],
) -> str:
    if not chat_history:
        return question

    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "Rewrite the user's latest question into a standalone question using the prior conversation. "
                "Resolve references like 'it', 'they', 'each', and 'those'. "
                "Keep the rewritten question grounded in the conversation and do not add new facts. "
                "Return only the rewritten standalone question.",
            ),
            (
                "human",
                "Conversation:\n{history}\n\nLatest question:\n{question}",
            ),
        ]
    )

    chain = prompt | llm | StrOutputParser()
    rewritten = chain.invoke(
        {
            "history": format_chat_history(chat_history),
            "question": question,
        }
    ).strip()
    return rewritten or question


def answer_question(
    llm: ChatOpenAI,
    question: str,
    retrieved_docs: list[Document],
    chat_history: list[dict[str, str]] | None = None,
) -> str:
    if not retrieved_docs:
        return "I could not find enough support for that answer in the uploaded document."

    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "You are a grounded document question-answering assistant. "
                "Use only the supplied document context. "
                "Never use outside knowledge. "
                "You may make simple, reasonable inferences when they are clearly supported by the context. "
                "Valid inference includes combining related statements, direct comparisons, and obvious conclusions from trends stated in the text. "
                "Do not speculate, invent missing details, or introduce background knowledge. "
                "If relevant information exists, provide the best grounded answer even when it requires combining multiple snippets. "
                "If the context does not contain enough relevant information, reply exactly: "
                "I could not find enough support for that answer in the uploaded document."
            ),
            (
                "human",
                "Conversation history:\n{history}\n\n"
                "Standalone question:\n{question}\n\n"
                "Context from the uploaded documents:\n{context}\n\n"
                "Return a clear and concise answer grounded only in the uploaded documents.",
            ),
        ]
    )

    chain = prompt | llm | StrOutputParser()
    return chain.invoke(
        {
            "question": question,
            "history": format_chat_history(chat_history or []),
            "context": format_context(retrieved_docs),
        }
    ).strip()


def summarize_pdf(
    llm: ChatOpenAI,
    chunked_documents: list[Document],
) -> str:
    if not chunked_documents:
        return "No readable document content was available for summary."

    map_prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "Summarize only the supplied document chunk. "
                "Do not add facts not present in the text.",
            ),
            (
                "human",
                "Summarize this chunk from the uploaded documents in 3 to 5 bullet-style sentences:\n\n{chunk}",
            ),
        ]
    )
    reduce_prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "Combine the partial summaries into one grounded document summary. "
                "Use only the supplied text. "
                "Structure the answer with Overview, Key Points, and Important Details.",
            ),
            (
                "human",
                "Partial summaries:\n{partial_summaries}",
            ),
        ]
    )

    map_chain = map_prompt | llm | StrOutputParser()
    partial_summaries = []

    batch_size = 4
    for start in range(0, len(chunked_documents), batch_size):
        batch = chunked_documents[start : start + batch_size]
        batch_text = "\n\n".join(doc.page_content for doc in batch)
        partial_summaries.append(map_chain.invoke({"chunk": batch_text}).strip())

    reduce_chain = reduce_prompt | llm | StrOutputParser()
    return reduce_chain.invoke({"partial_summaries": "\n\n".join(partial_summaries)}).strip()


def filter_documents_by_source(
    documents: list[Document],
    selected_source: str,
) -> list[Document]:
    if selected_source in {"All PDFs", "All Documents"}:
        return documents
    return [
        document
        for document in documents
        if str(document.metadata.get("source", "")) == selected_source
    ]


def build_citations(documents: Iterable[Document]) -> list[str]:
    citations = []
    for doc in documents:
        excerpt = _normalize_text(doc.page_content)[:220]
        citations.append(
            f"{doc.metadata.get('source', 'Unknown')} - {format_document_location(doc)}: {excerpt}"
        )
    return citations


def build_chunk_displays(
    chunks: list[RetrievedChunk],
    question: str,
    max_sentences: int = 2,
) -> list[ChunkDisplay]:
    terms = _keyword_terms(question)
    displays: list[ChunkDisplay] = []

    for chunk in chunks:
        sentences = _split_sentences(chunk.document.page_content)
        if sentences:
            ranked_sentences = sorted(
                sentences,
                key=lambda sentence: _keyword_overlap(question, sentence),
                reverse=True,
            )
            supporting = [sentence for sentence in ranked_sentences[:max_sentences] if sentence]
        else:
            fallback = re.sub(r"^[^\w(]+", "", chunk.document.page_content[:260].strip())
            supporting = [fallback] if fallback else []

        excerpt = " ".join(supporting).strip()
        if not excerpt:
            excerpt = _normalize_text(chunk.document.page_content)[:260]
        displays.append(
            ChunkDisplay(
                source=str(chunk.document.metadata.get("source", "Unknown")),
                location=format_document_location(chunk.document),
                relevance_score=round(chunk.combined_score, 2),
                highlighted_excerpt=_highlight_terms(excerpt, terms),
                supporting_sentences=[_highlight_terms(sentence, terms) for sentence in supporting],
            )
        )

    return displays


def compute_retrieval_metrics(
    retrieved_chunk_ids: list[str],
    relevant_chunk_ids: list[str],
    k: int,
) -> dict[str, float]:
    relevant_set = set(relevant_chunk_ids)
    retrieved_at_k = retrieved_chunk_ids[:k]
    hits = sum(1 for chunk_id in retrieved_at_k if chunk_id in relevant_set)

    precision = hits / k if k else 0.0
    recall = hits / len(relevant_set) if relevant_set else 0.0
    f1 = (2 * precision * recall / (precision + recall)) if (precision + recall) else 0.0

    return {
        "precision_at_k": round(precision, 4),
        "recall_at_k": round(recall, 4),
        "f1_at_k": round(f1, 4),
        "hits_at_k": hits,
    }
